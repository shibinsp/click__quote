import docx
from docx import Document
import json
import re

def extract_complete_document_content():
    """Extract all content from the Word document with detailed structure"""
    
    doc_path = r'quotaion template.docx'
    doc = Document(doc_path)
    
    print('=== COMPLETE DOCUMENT CONTENT EXTRACTION ===')
    print(f'Total paragraphs: {len(doc.paragraphs)}')
    print()
    
    # Extract all content with detailed structure
    all_content = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # Only include non-empty paragraphs
            # Get formatting information
            runs = paragraph.runs
            is_bold = any(run.bold for run in runs if run.bold)
            is_italic = any(run.italic for run in runs if run.italic)
            
            content_item = {
                'line': i + 1,
                'text': text,
                'is_bold': is_bold,
                'is_italic': is_italic,
                'length': len(text)
            }
            all_content.append(content_item)
    
    # Manual page breaks based on content analysis
    # Looking for natural page breaks in the document
    pages_content = {i: [] for i in range(1, 15)}  # 14 pages
    
    current_page = 1
    
    for item in all_content:
        text = item['text']
        
        # Page break indicators
        if any(indicator in text.lower() for indicator in [
            'page break', 'new page', '---', 'section break'
        ]):
            current_page += 1
            if current_page > 14:
                current_page = 14
            continue
            
        # Specific content-based page assignments
        if 'UK Power Networks (Operations) Ltd' in text and current_page == 1:
            pages_content[1].append(item)
        elif 'Quote valid until' in text.lower() or 'when can you expect' in text.lower():
            current_page = 2
            pages_content[2].append(item)
        elif 'please indicate which option you accept' in text.lower():
            current_page = 3
            pages_content[3].append(item)
        elif 'acceptance form' in text.lower():
            current_page = 4
            pages_content[4].append(item)
        elif 'scope of works' in text.lower():
            current_page = 5
            pages_content[5].append(item)
        elif 'your responsibilities' in text.lower():
            current_page = 8
            pages_content[8].append(item)
        elif 'payment profile' in text.lower():
            current_page = 11
            pages_content[11].append(item)
        else:
            # Distribute remaining content across pages
            if current_page <= 14:
                pages_content[current_page].append(item)
                
                # Move to next page after certain amount of content
                if len(pages_content[current_page]) > 15:  # Adjust based on content density
                    current_page += 1
                    if current_page > 14:
                        current_page = 14
    
    # Print detailed content for each page
    for page_num in range(1, 15):
        print(f'\n=== PAGE {page_num} CONTENT ===')
        page_items = pages_content[page_num]
        print(f'Lines on this page: {len(page_items)}')
        print('---')
        
        for item in page_items:
            formatting = []
            if item['is_bold']:
                formatting.append('BOLD')
            if item['is_italic']:
                formatting.append('ITALIC')
            
            format_str = f' [{", ".join(formatting)}]' if formatting else ''
            print(f'Line {item["line"]}: {item["text"]}{format_str}')
        
        print('---')
        if page_items:
            print(f'Total characters on page: {sum(item["length"] for item in page_items)}')
        else:
            print('No content on this page')
    
    print(f'\n=== SUMMARY ===')
    print(f'Total content lines: {len(all_content)}')
    print(f'Total characters: {sum(item["length"] for item in all_content)}')
    
    # Save structured data to JSON for easy processing
    structured_data = {
        'pages': pages_content,
        'total_lines': len(all_content),
        'total_characters': sum(item["length"] for item in all_content)
    }
    
    with open('docx_content_structured.json', 'w', encoding='utf-8') as f:
        json.dump(structured_data, f, indent=2, ensure_ascii=False)
    
    print('\nStructured data saved to docx_content_structured.json')
    
    return structured_data

if __name__ == "__main__":
    extract_complete_document_content()